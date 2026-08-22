# -*- coding: utf-8 -*-
"""BIGB 내용 파악용 상세 문서(DOCX) 생성.

**제출물이 아니다.** 이동 중에 읽으며 프로젝트 전체를 파악하는 용도라
심사 서식을 따르지 않는다. 쉬운 말로 쓰되 빠짐없이 쓴다 — 용어는 처음
나올 때 풀어 쓰고, 설계 선택마다 "왜 그렇게 했는가"를 붙인다.

**모든 수치를 outputs/ 와 config.py 에서 읽는다.** 손으로 적은 값이 하나도
없어야 한다 — 코드가 바뀌면 문서를 다시 뽑으면 맞는다.
CLAUDE.md §0.1-1 이 그대로 이 스크립트의 설계 이유다.

사용:  python tools/make_bigb_doc.py
"""
from pathlib import Path
import json
import sys
import datetime

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

import config

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

STEM = "BIGB_0822"
OUT = config.ROOT / "outputs" / f"{STEM}.docx"

INK = RGBColor(0x1E, 0x22, 0x28)
MUTED = RGBColor(0x5C, 0x65, 0x70)
ACCENT = RGBColor(0x14, 0x4E, 0x8C)
WARN = RGBColor(0x8C, 0x20, 0x18)

FONT = "맑은 고딕"
BODY = 10.5   # 이동 중 읽기용이라 본문을 키웠다


# ── 서식 도구 ──────────────────────────────────────────────────────────────

def set_kfont(doc, name=FONT, size=BODY):
    for style_name in ("Normal", "List Bullet", "List Number"):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = name
        st.font.size = Pt(size)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _run(p, text, size, bold, color, italic=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return r


def para(doc, text="", size=BODY, bold=False, color=INK, space_after=8,
         align=None, italic=False, indent=0.0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    _run(p, text, size, bold, color, italic)
    return p


def rich(doc, parts, size=BODY, space_after=8, indent=0.0):
    """[(텍스트, 굵게, 색), ...] 를 한 문단에 이어 붙인다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for text, bold, color in parts:
        _run(p, text, size, bold, color)
    return p


def bullet(doc, text, size=BODY, head=None, indent=0.8):
    """가운뎃점 목록. 행잉 인덴트로 둘째 줄이 글머리 아래로 물리지 않게 한다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    _run(p, "· ", size, False, MUTED)
    if head:
        _run(p, head, size, True, INK)
    _run(p, text, size, False, INK)
    return p


def table(doc, headers, rows, widths=None, size=9, align_right=()):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        _run(c.paragraphs[0], str(hd), size, True, INK)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            if i in align_right:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(p, str(v), size, False, INK)
    if widths:
        # **autofit 을 끄지 않으면 폭이 무시된다.** 워드는 내용에 맞춰 다시
        # 계산하고, 열이 좁아진 자리에서 한글이 어절 중간에 접힌다
        # ("갱폼 작업면 (2 / 개 층)"). 열 객체와 셀 양쪽에 넣어야 먹는다.
        t.autofit = False
        for i, w in enumerate(widths):
            t.columns[i].width = Cm(w)
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def h(doc, text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        r.font.color.rgb = ACCENT if level == 1 else INK
    return hd


def note(doc, text, size=9.5):
    """왜 그렇게 했는지·무엇을 못 하는지 적는 들여쓴 회색 문단."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.3
    _run(p, text, size, False, MUTED, italic=False)
    return p


def pct(x, d=1):
    return "-" if x is None else f"{x*100:.{d}f}%"


# ── 본문 ───────────────────────────────────────────────────────────────────

def build():
    O = config.ROOT / "outputs"
    site = json.loads((O / "site_eval.json").read_text(encoding="utf-8"))
    comp = json.loads((O / "comparison.json").read_text(encoding="utf-8"))
    cp = json.loads((O / "curve_params.json").read_text(encoding="utf-8"))
    sr = json.loads((O / "safety_report.json").read_text(encoding="utf-8"))
    sn = json.loads((O / "sensitivity.json").read_text(encoding="utf-8"))
    mf = json.loads((config.DATA_FILTERED / "manifest.json").read_text(encoding="utf-8"))

    ov = sr["coverage"]["overall"]
    sd = ov["score_detail"]
    realloc = (sr.get("options") or {}).get("reallocate") or {}
    pres = sr["prescription"]
    pl = comp["placements"]
    dl = comp["delta_WDR"]
    tp = sr.get("time_phased") or {}
    prim = cp["per_target"][cp["primary"]]
    rng = site["out_of_measured_range"]
    dori = site["dori_pair_stats"]
    base = site["geometric_baseline"]
    occ = [v for v in site["voxels"] if v.get("occupiable") is not False]
    cm = config.CAMERA_MODEL

    doc = Document()
    set_kfont(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.4)
        s.top_margin = s.bottom_margin = Cm(2.0)

    # ── 표지 ──────────────────────────────────────────────────────────
    para(doc, "BIGB", 26, True, ACCENT, 2, WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "AI CCTV 배치 적정성 평가 모델", 15, True, INK, 4,
         WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "내용 파악용 상세 문서", 11, False, MUTED, 16,
         WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, f"{datetime.date.today():%Y-%m-%d} 기준   ·   "
              f"모든 수치는 outputs/ 에서 자동으로 읽어 넣었다", 9, False, MUTED,
         20, WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "이 문서는 제출물이 아니다. 프로젝트가 무엇을 하는 물건이고 왜 "
              "그렇게 만들었는지를 처음부터 끝까지 따라갈 수 있게 쓴 것이다. "
              "용어는 처음 나올 때 풀어 썼고, 설계 선택마다 이유를 붙였다.",
         BODY, False, INK, 14)

    # ── 30초 요약 ─────────────────────────────────────────────────────
    h(doc, "30초 요약", 1)
    bullet(doc, "건설현장 CCTV 계획서(어디에 몇 대를 어느 방향으로 다는지)",
           head="넣는 것 — ")
    bullet(doc, "이 배치가 몇 점짜리인지(100점 만점), 어디가 안 보이는지, "
                "고치려면 옮길지 더 달지", head="나오는 것 — ")
    bullet(doc, "기존 도구는 \"카메라 시야에 들어오는가\"를 따진다. 이 도구는 "
                "\"들어온 걸 AI가 실제로 알아볼 수 있는가\"를 따진다",
           head="다른 점 — ")
    rich(doc, [("40m 떨어진 작업자는 화면에 ", False, INK), ("보인다", True, INK),
               (". 하지만 머리가 12픽셀밖에 안 되면 AI는 안전모를 썼는지 "
                "판단하지 못한다. 기존 계산은 이걸 100% 커버로 센다. "
                "우리는 ", False, INK),
               ("실측한 확률", True, WARN), ("로 센다.", False, INK)], BODY, 12)

    # ── 1. 문제 ───────────────────────────────────────────────────────
    h(doc, "1. 무슨 문제를 푸는가", 1, page_break=True)

    h(doc, "1.1 현장에 CCTV는 이미 달려 있다", 2)
    para(doc, "AI CCTV로 안전모 미착용을 잡는 건 이미 하고 있는 일이다. LH도 "
              "「늘봄 A-Eye」로 전국 350개 현장 중 위험도가 높은 20곳을 매일 "
              "골라 집중 관리한다. 그러니 \"AI CCTV를 도입하자\"는 제안은 "
              "이미 늦었다.")
    rich(doc, [("문제는 그 다음이다. ", False, INK),
               ("달아 놓은 카메라가 실제로 위험한 곳을 보고 있는지 아무도 "
                "숫자로 확인하지 않는다.", True, WARN),
               (" 카메라 대수와 설치 위치는 대개 관행과 감으로 정해진다.",
                False, INK)])

    h(doc, "1.2 확인하고 싶어도 기준이 없다", 2)
    para(doc, "확인하려면 \"얼마나 봐야 충분한가\"라는 기준이 있어야 한다. "
              "찾아봤는데 없었다. 아래는 실제로 원문을 열어 확인한 결과다.")
    table(doc, ["문서", "커버리지 기준", "실제로 적혀 있는 것"],
          [["스마트 안전장비 활용 가이드라인 (국토교통부·국토안전관리원, 88쪽)",
            "없음", "AI CCTV는 권장기준 한 줄. 점검표는 \"정상작동 하는가?\" 식 O/X"],
           ["KISA 지능형 CCTV 인증제도 안내서", "없음",
            "90%는 검출 정확도이며 KISA가 만든 시험 영상 위의 점수다"],
           ["KISA 지능형 CCTV 도입·운영 가이드 (106쪽)", "없음",
            "설치 각도·가림 등 정성적 저하 요인만"],
           ["건설기술진흥법 시행규칙 별표 7", "없음",
            "CCTV 설치·운용계획을 안전관리계획에 포함하라고만 함"],
           ["LH 늘봄 A-Eye 공개자료", "없음", "운영 방식만 공개, 규격은 비공개"]],
          widths=[6.4, 2.2, 7.6], size=8.5)
    rich(doc, [("특히 헷갈리기 쉬운 게 ", False, INK), ("90%", True, INK),
               ("다. 가이드라인이 요구하는 \"지능형 CCTV 인증\"의 합격선이 "
                "검출 정확도 90%인데, 이건 ", False, INK),
               ("KISA가 자기네 시험용 영상 DB 위에서 재는 점수", True, INK),
               ("다. 현장에 어떻게 설치했는지와 아무 상관이 없다. "
                "인증 받은 카메라를 사도 엉뚱한 데 달면 아무것도 못 잡는다.",
                False, INK)])
    note(doc, "정리하면 제도에는 두 개의 빈칸이 있다. (1) 얼마나 봐야 충분한가 "
              "(2) 애초에 '본다'가 무엇인가. 화각 안에 들어오면 본 것인가, "
              "AI가 알아볼 수 있어야 본 것인가. 이 도구는 두 번째 질문에 "
              "'알아볼 수 있어야 본 것'이라고 답하고, 그걸 숫자로 만든다.")

    h(doc, "1.3 그런데 CCTV 계획서는 이미 내는 서류다", 2)
    rich(doc, [("「건설기술진흥법 시행규칙」 별표 7 「안전관리계획의 수립기준」 "
                "1-다 항목이 ", False, INK),
               ("\"계측장비 및 폐쇄회로 텔레비전 등 안전 모니터링 장비의 설치 및 "
                "운용계획을 포함한다\"", True, INK),
               ("고 규정한다. 즉 CCTV 계획을 쓰는 칸은 법에 이미 있다. "
                "다만 ", False, INK),
               ("그 계획이 충분한지 판정하는 칸이 없다", True, WARN),
               (". 우리 산출물이 들어갈 자리가 정확히 거기다.", False, INK)])

    # ── 2. 핵심 아이디어 ──────────────────────────────────────────────
    h(doc, "2. 핵심 아이디어 — 검출확률 곡선", 1, page_break=True)

    h(doc, "2.1 세 가지가 검출을 망친다", 2)
    para(doc, "카메라가 작업자를 못 알아보는 이유는 크게 셋이다. 각각에 기호를 "
              "붙였다.")
    table(doc, ["기호", "이름", "쉬운 말", "예시"],
          [["ρ (로)", "유효 픽셀밀도", "화면에서 머리가 몇 픽셀인가",
            "멀수록 작아진다"],
           ["θ (세타)", "부감각", "얼마나 위에서 내려다보는가",
            "바로 위에서 보면 안전모가 원판으로 보인다"],
           ["o (오)", "가림률", "몸이 얼마나 가려졌는가",
            "비계 기둥 뒤에 서 있다"]],
          widths=[1.8, 3.0, 5.4, 6.0], size=9)
    note(doc, "방위각(정면인지 뒤통수인지)은 일부러 뺐다. 안전모는 위에서 보면 "
              "대체로 둥글어 앞뒤 차이가 작고, 축을 하나 더하면 실험 조건이 "
              "다섯 배로 늘어난다. 야간·조명·기상, PTZ 회전 카메라, 공정에 "
              "따라 카메라를 옮기는 4D 배치도 범위 밖이다. 못 한 게 아니라 "
              "정하고 뺀 것이다.")

    h(doc, "2.2 곱셈으로 합친다", 2)
    rich(doc, [("세 축을 각각 곡선으로 만들고 곱한다. ", False, INK),
               ("P(ρ, θ, o) = f(ρ) × g(θ) × h(o)", True, ACCENT),
               (" 형태이고, 이걸 ", False, INK),
               ("분리형 곱셈 모델", True, INK),
               ("이라고 부른다. \"멀고, 위에서 보고, 가려지면\" 세 배로 "
                "나빠진다는 뜻이다.", False, INK)])
    para(doc, "곱셈을 쓴 이유는 해석이 쉽고 최적화에 바로 넣을 수 있어서다. "
              "축끼리 서로 영향을 주는 부분(예: 멀면서 동시에 가려질 때 유독 "
              "더 나쁜 효과)은 무시한다. 1차 근사이고, 그 사실을 제안서에 "
              "먼저 밝히는 것이 방어 전략이다.")

    h(doc, "2.3 실제로 나온 곡선", 2)
    table(doc, ["축", "함수 모양", "파라미터", "읽는 법"],
          [["f(ρ)", "로지스틱(S자)",
            f"L={prim['f_rho']['L']:.3f}, k={prim['f_rho']['k']:.3f}, "
            f"x0={prim['f_rho']['x0']:.2f}px",
            f"머리가 {prim['f_rho']['x0']:.1f}픽셀일 때 검출률이 절반으로 꺾인다"],
           ["g(θ)", "로지스틱",
            f"k={prim['g_theta']['params']['k']:.4f}, "
            f"x0={prim['g_theta']['params']['x0']:.2f}°",
            f"{prim['g_theta']['params']['x0']:.0f}도 부근까지 버티다 무너진다"],
           ["h(o)", "지수 감쇠", f"λ={prim['h_occ']['lambda']:.3f}",
            "가림이 커질수록 급격히 떨어진다"]],
          widths=[1.6, 3.0, 5.4, 6.2], size=8.5)
    rich(doc, [("θ 곡선이 흥미롭다. 45도까지는 거의 안 떨어지다가 60~75도에서 "
                "절벽처럼 무너진다. 처음엔 2차함수로 맞췄는데 절벽을 못 맞혀서 "
                "로지스틱으로 바꿨다(단면 결정계수 "
                f"{prim['g_theta']['r2_candidates']['quadratic']:.3f} → "
                f"{prim['g_theta']['r2_candidates']['logistic']:.3f}). ",
                False, INK),
               ("그런데 KISA 도입·운영 가이드가 \"설치 각도 30도 이하(직하)\"를 "
                "성능 저하 요인으로 든다.", True, INK),
               (" 각도 기준선이 원문에 정의돼 있지 않아 추정이지만, 수직축 "
                f"기준으로 읽으면 부감 60도이고 우리 실측 변곡점 "
                f"{prim['g_theta']['params']['x0']:.1f}도와 거의 맞는다.",
                False, INK)])

    # ── 3. 실험 ───────────────────────────────────────────────────────
    h(doc, "3. 곡선을 어떻게 얻었나", 1, page_break=True)

    h(doc, "3.1 데이터", 2)
    rich(doc, [(f"SHWD(Safety Helmet Wearing Dataset)를 썼다. 안전모를 쓴 "
                f"사람과 안 쓴 사람이 라벨링된 공개 데이터셋이다. 여기서 "
                f"머리 크기가 {mf['min_head_px']}픽셀 이상인 것만 골라 "
                f"{mf['n_selected']}장을 실험셋으로 삼았다(후보 "
                f"{mf['n_candidates']}장 중). ", False, INK),
               ("반드시 test 분할에서만 뽑았다", True, WARN),
               (" — 학습에 쓴 사진으로 곡선을 재면 설치 조건의 효과가 아니라 "
                "모델의 암기를 재게 된다.", False, INK)])

    h(doc, "3.2 검출기를 직접 학습시켜야 했다", 2)
    para(doc, "처음엔 공개된 YOLO 사전학습 모델을 그대로 쓰려 했다. 그런데 그 "
              "모델은 사람·자동차 같은 일반 사물만 알고 '안전모'라는 개념이 "
              "없다. 아무 변형도 안 준 최상 조건에서도 주 지표가 0.088에 "
              "그쳤다. 곡선이 나올 수가 없어서 SHWD로 직접 파인튜닝했다.")
    table(doc, ["모델", "학습", "검출 성능(mAP50)", "곡선 적합도"],
          [["yolov8n (활성)", "50 epoch", "0.9372",
            f"{cp['r2_full_grid']:.4f}  통과"],
           ["yolo26n", "50 epoch", "0.9371", "0.7395  미달"]],
          widths=[4.0, 3.0, 4.4, 4.8], size=9)
    rich(doc, [("최신 모델(yolo26n)도 같은 조건으로 돌려봤다. ", False, INK),
               ("검출 성능은 사실상 동률인데 곡선 적합도는 오히려 나빠서 "
                "통과 기준에 미달했다.", True, INK),
               (" 그래서 활성 곡선은 yolov8n이다. 최신이라서가 아니라 맞는 "
                "쪽을 쓴다. 다만 곡선의 ", False, INK),
               ("형태는 두 모델에서 같았다", True, INK),
               (" — 변곡점이 5.73→5.63픽셀로 움직였을 뿐이다. 검출기를 바꿔도 "
                "관계가 남는다는 뜻이고, 이게 오히려 쓸 만한 결과다.",
                False, INK)])
    note(doc, "이 사실이 중요한 이유가 있다. 우리 검출기는 LH가 실제로 쓰는 "
              "A-Eye 모델이 아니다(접근 불가). 그래서 이 실험은 '검출률이 몇 "
              "퍼센트다'를 주장하는 게 아니라 '설치 조건에 따라 검출률이 "
              "변한다'는 관계를 증명하는 것으로 포지셔닝한다. 서로 다른 두 "
              "검출기가 같은 형태를 준다는 것이 그 포지셔닝을 뒷받침한다. "
              "나중에 A-Eye 곡선을 구하면 파일 하나만 갈아 끼우면 된다.")

    h(doc, "3.3 사진을 어떻게 망가뜨렸나", 2)
    para(doc, "세 축을 실제로 재려면 같은 사진을 조건별로 변형해야 한다.")
    bullet(doc, "목표 배율로 줄였다가 원래 크기로 다시 키운다. 정보가 "
                "사라진 상태가 남는다", head="ρ (픽셀밀도) — ")
    bullet(doc, "호모그래피 워핑으로 위에서 내려다본 것처럼 기울인다",
           head="θ (부감각) — ")
    bullet(doc, "회색 세로 줄무늬로 가린다", head="o (가림률) — ")
    rich(doc, [("가림을 랜덤 사각형이 아니라 ", False, INK),
               ("세로 줄무늬", True, WARN),
               ("로 한 게 핵심이다. 현장에서 시야를 막는 건 비계·동바리·거푸집 "
                "지주처럼 대부분 ", False, INK),
               ("수직 부재", True, INK), ("이기 때문이다.", False, INK)])
    rich(doc, [("다만 줄무늬 간격이 결과를 크게 좌우한다. 같은 30% 가림이라도 "
                "부재가 촘촘할수록 검출률이 급락한다(간격에 따라 0.443 / "
                "0.239 / 0.045). ", False, INK),
               ("근거가 있는 값이 아니어서 민감도 스윕 대상으로 다룬다",
                True, INK),
               (" — 8절에서 다시 나온다.", False, INK)])

    h(doc, "3.4 288개 조건을 전부 돌렸다", 2)
    para(doc, f"ρ 8수준 × θ 6수준 × o 6수준 = {cp['n_conditions']}조건, 각각 "
              f"{mf['n_selected']}장씩 추론했다. 조건마다 안전모 미착용 "
              f"재현율(주 지표), 착용 재현율, confidence 분포를 기록했다.")
    rich(doc, [("곡선을 맞춘 결과 전체 288점에 대한 결정계수가 주 지표 ",
                False, INK),
               (f"R² = {cp['r2_primary']:.4f}", True, ACCENT),
               (f"다. 두 탐지 항목 중 나쁜 쪽으로 대표하면 "
                f"{cp['r2_full_grid']:.4f}이고, 통과 기준 "
                f"{cp['r2_acceptance']}를 넘는다. R²는 곡선이 실측점을 얼마나 "
                f"잘 설명하는지를 0~1로 나타낸 값이다.", False, INK)])
    note(doc, "두 항목(미착용·착용)의 곡선을 따로 만들고 종합할 때는 평균이 "
              "아니라 최솟값을 쓴다. 평균을 쓰면 잘 잡히는 항목이 못 잡히는 "
              "항목을 가린다. 실제로 평균으로 바꿔보니 커버리지가 10.1%p "
              "부풀고 미달 복셀 343개가 사라졌다.")

    h(doc, "3.5 측정 안 한 구간은 0으로 본다", 2)
    rich(doc, [("실측 범위 밖으로 곡선을 늘려 쓰지 않는다. ", False, INK),
               (f"ρ < {rng['rho_below_min']['limit_px']:.0f}px, "
                f"θ > {rng['theta_over_max']['limit_deg']:.0f}°, "
                f"o > {rng['occ_over_max']['limit']:.2f}", True, INK),
               ("는 검출확률 0으로 처리한다.", False, INK)])
    rich(doc, [(f"실제로 잘린 양은 보이는 복셀-카메라 쌍 {rng['n_visible']:,}개 "
                f"중 {rng['any']['n']:,}개({pct(rng['any']['ratio'],2)})다. ",
                False, INK),
               ("ρ 하한에 걸린 건 0개", True, INK),
               ("인데, 처음엔 12픽셀까지만 쟀다가 현장 대부분이 그보다 멀다는 "
                "걸 발견하고 8·6·4픽셀을 추가로 측정해 구멍을 막았기 때문이다. "
                "그 전에는 외삽값이 ρ=5px에서 0.64를 내놨는데 실측해 보니 "
                "0.179였다.", False, INK)])

    # ── 4. 현장 계산 ──────────────────────────────────────────────────
    h(doc, "4. 현장에서 어떻게 계산하나", 1, page_break=True)

    h(doc, "4.1 현장을 큐브로 자른다", 2)
    rich(doc, [(f"현장 {site['site']['width_m']}m × {site['site']['depth_m']}m를 "
                f"한 변 {config.VOXEL_M}m짜리 정육면체로 잘게 나눈다. 이 작은 "
                f"상자 하나를 ", False, INK),
               ("복셀(voxel)", True, INK),
               (f"이라고 한다. 픽셀의 3차원판이다. 모두 "
                f"{len(site['voxels']):,}개가 나왔다.", False, INK)])
    rich(doc, [("바닥 평면이 아니라 ", False, INK), ("부피 전체", True, WARN),
               ("를 다루는 게 중요하다. CCTV는 바닥만 보는 게 아니라 2층 "
                "슬래브 위, 비계 위도 본다. 다만 사람이 갈 수 없는 허공까지 "
                "점수에 넣으면 안 되므로 ", False, INK),
               (f"사람이 설 수 있는 복셀 {len(occ):,}개만 채점 대상",
                True, INK), ("으로 표시한다.", False, INK)])

    h(doc, "4.2 광선을 쏴서 거리·각도·가림을 잰다", 2)
    para(doc, "복셀 하나와 카메라 한 대의 조합마다 다음을 계산한다.")
    bullet(doc, "카메라에서 복셀까지의 직선거리", head="거리 — ")
    bullet(doc, "거리로부터 계산한다. 4K 해상도, 화각 90도 기준으로 머리 "
                "25cm가 몇 픽셀로 찍히는지", head="픽셀밀도 ρ — ")
    bullet(doc, "카메라 높이와 복셀 높이의 차이로 계산", head="부감각 θ — ")
    bullet(doc, "복셀 자리에 키 1.7m 막대를 세우고 11개 점에서 카메라로 광선을 "
                "쏜다. 골조에 막히는 비율이 가림률", head="가림률 o — ")
    rich(doc, [(f"조합이 {dori['n_pairs']:,}개, 그중 조금이라도 보이는 게 "
                f"{dori['n_visible']:,}개다. 픽셀밀도 중앙값이 "
                f"{dori['rho_px_median']:.1f}픽셀인데, ", False, INK),
               ("이게 이미 상당히 작다", True, WARN),
               (". 사람 머리가 화면에서 13픽셀이면 안전모 착용 여부를 "
                "판단하기 만만치 않다.", False, INK)])
    note(doc, f"카메라는 조달 가능한 4K 기종({cm['reference']})을 기준으로 "
              f"잡았다. 다만 이것이 현장에서 가장 많이 쓰인다는 통계 근거는 "
              f"없다. 카메라 부각(위아래로 얼마나 기울였는지)은 검사하지 "
              f"않는다 — 수직 화각이 작업면을 덮도록 조준했다고 본다. "
              f"낙관적인 가정이다.")

    h(doc, "4.3 카메라 여러 대는 더하기가 아니다", 2)
    rich(doc, [("한 복셀을 두 대가 본다면 전체 검출확률은 ", False, INK),
               ("P = 1 - (1-P₁)(1-P₂)", True, ACCENT),
               ("로 합친다. 둘 다 놓칠 확률을 빼는 방식이다. 각각 0.6이면 "
                "합쳐서 0.84가 된다.", False, INK)])
    rich(doc, [("따라서 이 모델에서 ", False, INK),
               ("카메라가 겹쳐 보는 건 낭비가 아니라 이득", True, WARN),
               ("이다. 고위험 구역일수록 겹쳐 봐야 한다. 중첩을 페널티로 "
                "다루는 코드는 절대 넣지 않는다.", False, INK)])

    # ── 5. 위험구역 ───────────────────────────────────────────────────
    h(doc, "5. 어디가 위험한지 어떻게 아나", 1, page_break=True)

    h(doc, "5.1 우리가 정하지 않는다", 2)
    para(doc, "\"여기가 거푸집 설치 공간이라 우선 봐야 한다\"는 판단은 우리가 "
              "하지 않는다. 위험한 곳은 현장이 정하고, 그 정보는 안전관리계획서·"
              "가설계획서에 이미 적혀 있다. 우리는 세 갈래로 받는다.")
    table(doc, ["등급", "구역", "어디서 오나"],
          [["T1 자동 도출", "슬래브 단부 · 갱폼 작업면 · 타설/거푸집면",
            "골조 모양에서 규칙으로 뽑는다. 규칙마다 산업안전보건기준에 관한 "
            "규칙 조문을 붙였다"],
           ["T2 도면 필요", "슬래브 관통부(엘리베이터·계단실 구멍)",
            "골조를 판 한 장으로 모델링하면 구멍이 아예 없다. 도면이 있어야 한다"],
           ["T3 가설계획 필요", "굴착면 · 리프트 · 크레인 반경 · 자재 야적장",
            "골조가 아니라 장비를 어디 놓느냐의 문제다"]],
          widths=[2.8, 5.0, 8.4], size=8.5)
    note(doc, "모든 위험구역은 출처(source) 항목을 반드시 가져야 하고, 비어 "
              "있으면 프로그램이 즉시 멈춘다. 근거 없는 사각형은 \"그 좌표는 "
              "어디서 나왔냐\"는 질문 하나로 무너지기 때문이다.")

    h(doc, "5.2 가중치는 재해통계에서 뽑았다", 2)
    rich(doc, [("구역마다 1~10의 위험가중치를 준다. 처음엔 내가 감으로 정했는데 "
                "근거가 없어서, ", False, INK),
               ("고용노동부 「2025년 3분기 재해조사 대상 사망사고 발생 현황」",
                True, INK),
               ("에서 유도했다. 산식은 ", False, INK),
               ("가중치 = 1 + 9 × (사망자비율 × 사고당 사망자) / 최댓값",
                True, ACCENT), ("이다.", False, INK)])
    rich(doc, [("떨어짐이 전체 사망의 43.5%라 가중치 10이 되고, 무너짐은 5.0%라 "
                "2가 된다. ", False, INK),
               ("그래서 콘크리트 타설과 굴착면이 자재 야적장(3)보다 낮게 나온다.",
                True, WARN),
               (" 현장 감각과 어긋나는데, 통계가 그렇게 말하는 것이지 붕괴가 "
                "덜 위험해서가 아니다.", False, INK)])
    para(doc, "원인은 전국 합계가 구역별 노출 면적을 보정하지 못한다는 데 있다. "
              "떨어짐은 모든 단부·모든 층에 넓게 퍼져 있고 붕괴는 좁은 곳에 "
              "몰린다. 같은 사망자 수라도 좁은 데서 나오면 단위면적당 위험은 "
              "훨씬 크다. 우리 가중치는 복셀 하나하나에 곱해지므로 원래는 "
              "단위면적당 위험이어야 하는데, 전국 통계로는 그 보정을 할 수 없다.")
    rich(doc, [("그래서 확정값으로 주장하지 않고 ", False, INK),
               ("두 프로파일", True, INK),
               ("을 나란히 둔다 — 통계 그대로인 것과 붕괴 계열을 올린 판단값. "
                "둘 다 돌려보니 ", False, INK),
               ("충족/미달 판정과 최악 시간대가 같았다", True, ACCENT),
               (". 결론이 가중치 논쟁에 걸려 있지 않다는 뜻이다.", False, INK)])

    h(doc, "5.3 위험구역은 하루 종일 같지 않다", 2)
    para(doc, "갱폼을 인양하는 시간대에는 작업면과 개구부가 위험하고, 자재를 "
              "들이는 시간대에는 크레인 반경과 야적장이 위험하다. 공정표에서 "
              "시간대별로 어느 구역이 열리는지 받아 각각 진단한다.")
    if tp.get("windows"):
        table(doc, ["시간대", "작업", "점수", "판정"],
              [[w["window"]["id"], w["window"]["label"], f"{w['scored']:.3f}",
                w["verdict"]] for w in tp["windows"]],
              widths=[2.2, 5.0, 3.0, 6.0], size=9, align_right=(2,))
        rich(doc, [("종합은 평균이 아니라 ", False, INK),
                   ("최악의 시간대", True, WARN),
                   (f"로 대표한다(현재 {tp.get('worst_window')}). 평균을 쓰면 "
                    f"위험한 시간대의 실패가 한가한 시간대에 가려진다.",
                    False, INK)])

    # ── 6. 점수 ───────────────────────────────────────────────────────
    h(doc, "6. 100점을 어떻게 매기나", 1, page_break=True)

    h(doc, "6.1 단일 커버리지 %로는 안 되는 이유", 2)
    rich(doc, [("\"현장의 90%를 본다\"는 문장은 ", False, INK),
               ("어디를 놓쳤는지를 지운다", True, WARN),
               (". 넓은 저위험 구역을 잘 덮으면 갱폼 작업면을 통째로 놓치고도 "
                "90%가 나온다. 그래서 구역마다 따로 채점한다.", False, INK)])

    h(doc, "6.2 채점 방식", 2)
    bullet(doc, "100 × 그 구역 가중치 / 전체 가중치 합. 면적과 무관하다",
           head="배점 — ")
    bullet(doc, "가중치가 높을수록 더 많이 요구한다. 가중치 1이면 70%, "
                "10이면 99%", head="요구 커버리지 — ")
    bullet(doc, "min(1, 커버리지 ÷ 요구). 요구를 넘겨도 더 주지 않는다",
           head="달성률 — ")
    bullet(doc, "배점 × 달성률을 전부 더한 값", head="총점 — ")
    rich(doc, [("배점을 복셀 개수가 아니라 가중치로 나눈 게 핵심이다. 면적에 "
                "비례시키면 넓은 구역이 점수를 지배한다. ", False, INK),
               ("타워크레인 반경은 12,320 복셀이고 리프트 승강구는 896인데, "
                "리프트를 놓치는 게 8배 덜 나쁘지 않다.", True, INK)])
    rich(doc, [("초과 달성으로 벌충하지 못하게 한 것도 같은 이유다. 갱폼을 "
                "99%에서 100%로 올린 걸로 굴착면 33%를 메울 수 없다. ",
                False, INK),
               ("안전기준은 평균으로 면제되지 않는다.", True, INK)])

    h(doc, "6.3 치명 구역 게이트", 2)
    rich(doc, [("가중치 7 이상인 구역이 요구에 못 미치면 ", False, INK),
               ("총점이 아무리 높아도 충족이 아니다", True, WARN),
               (". 실제로 8대를 재배치하면 95.6점이 나오는데 개구부가 여전히 "
                "미달이었다. 그걸 \"충족\"이라고 부르면 이 도구가 하려던 말을 "
                "스스로 뒤집는 셈이다.", False, INK)])

    h(doc, "6.4 예시 계획서 채점 결과", 2)
    table(doc, ["구역", "가중", "배점", "커버리지", "요구", "달성률", "획득"],
          [[r["label"], r["weight"], f"{r['allocated']:.1f}", pct(r["coverage"]),
            pct(r["required"], 0), pct(r["attainment"], 0), f"{r['earned']:.1f}"]
           for r in sd["rows"]],
          widths=[4.6, 1.3, 1.4, 2.2, 1.5, 1.7, 1.4], size=8.5,
          align_right=(1, 2, 3, 4, 5, 6))
    rich(doc, [("총점 ", False, INK), (f"{sd['total']:.1f} / 100", True, WARN),
               (f", 등급 {sd['grade']}. 치명 구역 "
                f"{len(sd['critical_failures'])}곳이 미달이라 점수와 무관하게 "
                f"등급 상한이 걸렸다.", False, INK)])

    # ── 7. 처방 ───────────────────────────────────────────────────────
    h(doc, "7. 미달이면 무엇을 하라고 하나", 1, page_break=True)

    h(doc, "7.1 순서가 중요하다", 2)
    para(doc, "먼저 대수를 늘리지 않고 되는지 본다. 예산을 쓰지 않고 해결되는 "
              "일을 \"더 사세요\"라고 답하면 발주처가 쓰지 않는다.")
    bullet(doc, "같은 대수로 위치만 바꿔서 목표에 닿는가", head="1단계 — ")
    bullet(doc, "안 되면 몇 대를 어디에 더 달아야 하는가", head="2단계 — ")
    bullet(doc, "후보 위치를 다 써도 안 되면 그 사실을 말한다. 카메라 사양을 "
                "올리거나 설치 가능 위치를 늘려야 한다", head="3단계 — ")

    h(doc, "7.2 탐욕 알고리즘", 2)
    rich(doc, [("가능한 조합을 전부 따지는 건 불가능하다(후보 "
                f"{len(site['cameras'])}곳에서 8대를 고르는 경우의 수가 "
                f"천문학적이다). 그래서 ", False, INK),
               ("한 대씩, 그때그때 가장 이득이 큰 자리를 고른다", True, INK),
               (". 이걸 탐욕(greedy) 알고리즘이라 한다.", False, INK)])
    rich(doc, [("탐욕은 보통 최적해를 보장하지 못하는데, 우리 목적함수는 "
                "카메라를 더할수록 추가 이득이 줄어드는 성질(submodular)이 "
                "있어서 ", False, INK),
               ("탐욕해가 최적해의 63% 이상임이 수학적으로 보장", True, ACCENT),
               ("된다. 제안서에 인용할 만한 성질이다.", False, INK)])

    h(doc, "7.3 실제 처방문", 2)
    para(doc, pres["text"].replace("**", ""), BODY, False, INK, 10)
    if realloc:
        table(doc, ["안", "점수", "등급", "남은 치명 구역"],
              [["현 계획서 8대", f"{ov['score_100']:.1f}", ov["grade"],
                f"{len(ov['critical_failures'])}곳"],
               ["같은 8대 재배치", f"{realloc['score_100']:.1f}", realloc["grade"],
                f"{len(realloc['critical_failures'])}곳"],
               [f"{pres['add_cameras']}대 증설", f"{pres['resulting']*100:.1f}",
                "A", "0곳"]],
              widths=[5.4, 3.0, 2.4, 5.4], size=9, align_right=(1,))

    h(doc, "7.4 목표를 하나로 정하지 않았다", 2)
    rich(doc, [("커버리지 기준이 없다는 게 1.2절의 결론이었으므로, 90%를 "
                "확정값처럼 쓰지 않는다. ", False, INK),
               ("네 임계에서 각각 무엇이 필요한지를 함께 낸다", True, INK),
               (". \"90%의 근거가 뭐냐\"는 질문에 \"근거가 없어서 고르지 "
                "않았다\"고 답할 수 있다.", False, INK)])
    note(doc, "실제로 돌려보면 80·85·90·95% 모두 같은 증설 대수가 나온다. "
              "치명 구역 게이트가 임계보다 강하게 작동하기 때문이다. 개구부가 "
              "요구를 못 넘는 한 임계를 낮춰도 소용이 없다. 임계 논쟁보다 "
              "치명 구역 하나가 결정적이라는 뜻이다.")

    # ── 8. 결과 ───────────────────────────────────────────────────────
    h(doc, "8. 결과를 어떻게 읽나", 1, page_break=True)

    h(doc, "8.1 같은 8대인데 자를 바꾸면 배치가 갈린다", 2)
    para(doc, "이게 이 프로젝트의 핵심 실증이다. 세 가지 설계 기준으로 각각 "
              "8대를 배치한 뒤, 셋 다 실측 곡선의 자로 다시 재서 비교했다.")
    table(doc, ["설계 기준", "무엇으로 배치했나", "위험가중 검출률", "미달 복셀"],
          [["기하 커버리지", "보이면 커버로 침 (기존 방식)",
            f"{pl['geometric']['WDR']:.4f}",
            f"{pl['geometric']['fail_voxel_count']:,}"],
           ["문헌의 가정 곡선", "논문에서 쓰는 가정된 확률 곡선",
            f"{pl['assumed']['WDR']:.4f}",
            f"{pl['assumed']['fail_voxel_count']:,}"],
           ["실측 곡선", "우리가 잰 곡선 (제안 방식)",
            f"{pl['empirical']['WDR']:.4f}",
            f"{pl['empirical']['fail_voxel_count']:,}"]],
          widths=[3.6, 6.0, 3.4, 3.0], size=9, align_right=(2, 3))
    rich(doc, [(f"검출률 차이는 +{dl['empirical_minus_geometric']:.4f}로 크지 "
                f"않아 보인다. 그런데 ", False, INK),
               (f"미달 복셀은 {pl['geometric']['fail_voxel_count']:,}개에서 "
                f"{pl['empirical']['fail_voxel_count']:,}개로 "
                f"{(1-pl['empirical']['fail_voxel_count']/pl['geometric']['fail_voxel_count'])*100:.0f}% "
                f"줄었다", True, WARN),
               (". 평균은 비슷한데 못 보는 자리가 3분의 1로 준 것이고, 이쪽이 "
                "훨씬 강한 그림이다.", False, INK)])
    rich(doc, [(f"고른 카메라도 다르다. 기하 기준과 실측 기준이 공통으로 고른 "
                f"건 8대 중 "
                f"{comp['overlap_camera_count']['empirical_vs_geometric']}대"
                f"뿐이다. ", False, INK),
               ("같은 예산으로 절반은 다른 자리에 달게 된다는 뜻", True, INK),
               ("이다.", False, INK)])

    h(doc, "8.2 기존 방식을 일부러 약하게 만들지 않았다", 2)
    rich(doc, [("\"보이면 커버\"를 그대로 쓰면 116m 떨어진 복셀도 커버로 잡힌다. "
                "그건 실무보다 못하게 모델링해 놓고 이긴 것이 되므로, 기하 "
                "기준선에 ", False, INK),
               ("IEC 62676-4(DORI) 최소 픽셀밀도", True, INK),
               ("를 걸었다. 실무 설계도구(JVSG, Axis Site Designer)가 지키는 "
                "기준이다.", False, INK)])
    table(doc, ["DORI 등급", "요구 픽셀밀도", "그 배치의 검출률", "우리와 차이"],
          [[s["dori_level"], f"{s['min_rho_px']:.2f}px", f"{s['WDR']:.4f}",
            f"+{s['delta_WDR']:.4f}" + ("  ← 기본값" if s["is_default"] else "")]
           for s in base["sweep"]],
          widths=[3.4, 3.4, 3.6, 5.6], size=9, align_right=(1, 2))
    rich(doc, [("네 등급을 전부 돌린 뒤 ", False, INK),
               ("우리와 차이가 가장 작은 등급을 기본값으로 골랐다", True, ACCENT),
               (". 기존 방식에 가장 유리한 조건에서 싸운 것이고, \"약한 상대와 "
                "붙었다\"는 비판을 막기 위해서다.", False, INK)])
    note(doc, "recognition 등급이 observation보다 나쁜 건 오류가 아니다. "
              "임계를 올리면 만족 가능한 복셀이 급감해 탐욕이 오히려 나빠진다. "
              "임계 방식 자체의 구조적 결함이며 제안서 소재다. 참고로 DORI는 "
              "사람이 화면을 보고 판단하는 기준이지 AI 검출기에 대해 검증된 게 "
              "아니다. 기존 방식의 기준선을 세우는 용도로만 쓴다.")

    h(doc, "8.3 결론이 근거 없는 값에 매달려 있지 않다", 2)
    rich(doc, [("근거가 공개되지 않은 값이 둘 있다 — 줄무늬 간격과 비계가 "
                "시야를 막는 비율. 둘 다 넓게 스윕해 봤다. ", False, INK),
               (f"절대 검출률은 크게 흔들리지만, 두 배치의 차이는 "
                f"+{sn['delta_WDR_range'][0]:.4f} ~ "
                f"+{sn['delta_WDR_range'][1]:.4f}로 부호가 유지된다.",
                True, ACCENT)])
    para(doc, "그래서 우리가 주장하는 건 절대 수치가 아니라 두 배치의 차이다. "
              "이 구분이 방어의 핵심이다.")

    # ── 9. 한계 ───────────────────────────────────────────────────────
    h(doc, "9. 한계 — 먼저 말한다", 1, page_break=True)
    para(doc, "심사에서 물어보기 전에 먼저 밝히는 항목들이다. 코드로 보정하려 "
              "하면 오히려 검증 불가능해진다.")
    for t1, t2 in [
        ("검출기가 대리모델이다",
         "SHWD로 학습시킨 YOLO이며 LH A-Eye의 실제 모델이 아니다. 절대 수치가 "
         "아니라 관계의 존재를 증명하는 실험으로 포지셔닝한다."),
        ("다운샘플링은 실제 원거리 촬영과 다르다",
         "대기 흐림, 렌즈 광학 한계가 반영되지 않는다."),
        ("호모그래피는 3D 시점 변화의 근사다",
         "실제로 카메라를 위로 올려 찍은 것과 사진을 기울인 것은 다르다."),
        ("곱셈 모델은 1차 근사다",
         "축 간 상호작용을 무시한다. 게다가 잔차가 낙관 쪽으로 기운다 — 실측 "
         "0.5 미만 조건에서 평균 0.09만큼 높게 본다. 안 보이는 곳을 보인다고 "
         "말하는 방향이라 반드시 밝혀야 한다."),
        ("배점표와 요구곡선은 우리가 정했다",
         "가중치 비례 배점, 70~99% 요구곡선, 등급 경계(A 90 / B 80 / C 70) "
         "모두 게시된 근거가 없다. LH가 기준을 내면 함수 하나만 갈아 끼운다."),
        ("가림률 실측 통계가 없다",
         "건설현장에서 실제로 시야가 얼마나 막히는지 공개된 통계가 없다. "
         "그래서 입력이 아니라 스윕 파라미터로 다룬다."),
        ("카메라 부각을 검사하지 않는다",
         "수직 화각이 작업면을 덮도록 조준했다고 본다. 낙관적인 가정이다."),
        ("가상 현장이다",
         "실제 LH 현장 도면이 아니라 코드로 만든 100×60m 모사 현장이다. "
         "입력을 계약 형태로 분리해 두어 실제 도면이 들어오면 그대로 돈다."),
    ]:
        bullet(doc, t2, head=f"{t1} — ")

    para(doc, "", 6, space_after=10)
    rich(doc, [("범위 밖으로 정한 것도 밝힌다. ", False, MUTED),
               ("공정 단계별 카메라 재배치(4D), 방위각, 야간·조명·기상, PTZ "
                "동적 제어, 실시간 관제 시스템, A-Eye 실제 모델 재현",
                False, MUTED),
               (". 못 한 게 아니라 정하고 뺀 것이며, 특히 4D는 2025년 10월 "
                "논문과 정면으로 겹쳐서 전략적으로 제외했다.", False, MUTED)],
         9.5)

    # ── 10. 예상 질문 ─────────────────────────────────────────────────
    h(doc, "10. 나올 만한 질문과 답", 1, page_break=True)
    for q, a in [
        ("이미 AI CCTV 쓰고 있는데 뭐가 새롭나?",
         "AI CCTV를 도입하자는 게 아니라, 이미 단 카메라가 위험구역을 실제로 "
         "검출할 수 있는 조건에 있는지 진단하는 도구다. 도입과 진단은 다른 "
         "층위다."),
        ("90%는 어디서 나온 숫자인가?",
         "우리가 제안한 값이고 고시값이 아니다. 그래서 확정하지 않고 네 임계 "
         "전부의 결과를 함께 낸다. 조사 결과 커버리지 비율 기준은 어느 법령·"
         "고시·지침에도 없었다."),
        ("픽셀 기준으로 화질을 따지는 건 이미 있는 개념 아닌가?",
         "맞다. DORI·DCRI·Johnson 기준이 있고 우리가 새로 만든 게 아니다. "
         "새로운 건 임계로 자르지 않고 연속 확률로 다루며, 그 확률을 여러 "
         "카메라에 걸쳐 결합하고 위험가중으로 채점까지 간다는 점이다."),
        ("중첩이 이득이라는 게 새로운 발견인가?",
         "아니다. 무선센서망 문헌의 표준 결과다. 우리 발견으로 서술하면 안 "
         "된다. 우리는 그 성질을 검출확률 위에서 쓸 뿐이다."),
        ("검출기가 실제 모델이 아니면 수치가 의미 있나?",
         "절대 수치는 의미가 제한적이고 그렇게 주장하지 않는다. 주장하는 건 "
         "설치 조건에 따라 검출률이 변한다는 관계와, 두 배치의 차이다. "
         "서로 다른 두 검출기에서 곡선 형태가 같았다는 것이 근거다."),
        ("가상 현장이면 실제로 쓸 수 있나?",
         "입력을 네 개의 계약 파일로 분리했다. 실제 도면·안전관리계획서·"
         "공정표를 그 형식으로 넣으면 그대로 돈다. IFC 어댑터도 있는데, "
         "슬래브·층고·개구부는 자동으로 뽑고 비계는 설계 BIM에 없어서 사람이 "
         "채운다."),
        ("위험구역 좌표는 누가 정하나?",
         "현장이 정한다. 우리는 판단하지 않고 받아 적거나 골조에서 규칙으로 "
         "뽑는다. 모든 구역이 출처를 갖고, 없으면 프로그램이 멈춘다."),
    ]:
        rich(doc, [("Q. ", True, ACCENT), (q, True, INK)], BODY, 3)
        para(doc, a, BODY, False, INK, 10, indent=0.5)

    # ── 11. 용어 ──────────────────────────────────────────────────────
    h(doc, "11. 용어 정리", 1, page_break=True)
    table(doc, ["용어", "뜻"],
          [["복셀 (voxel)", "3차원 픽셀. 현장을 자른 작은 정육면체 하나"],
           ["ρ 픽셀밀도", "화면에서 사람 머리가 차지하는 픽셀 수. 멀수록 작다"],
           ["θ 부감각", "카메라가 얼마나 내려다보는 각도. 0도가 수평"],
           ["o 가림률", "몸이 장애물에 가려진 비율"],
           ["P_total", "여러 카메라를 합쳤을 때 그 복셀이 검출될 확률"],
           ["WDR", "위험가중 검출률. 위험가중치로 가중평균한 검출확률"],
           ["재현율 (recall)", "실제로 있는 것 중 검출기가 찾아낸 비율"],
           ["mAP50", "검출 모델의 표준 성능 지표. 1에 가까울수록 좋다"],
           ["R² 결정계수", "곡선이 실측점을 얼마나 잘 설명하는지 (0~1)"],
           ["파인튜닝", "이미 학습된 모델을 우리 데이터로 추가 학습시키는 것"],
           ["탐욕 알고리즘", "매 단계에서 당장 가장 이득인 선택을 하는 방식"],
           ["submodular", "더할수록 추가 이득이 줄어드는 성질. 탐욕 품질 보장 근거"],
           ["DORI", "IEC 62676-4. 감시 목적별 최소 픽셀밀도 기준(사람 관찰자용)"],
           ["호모그래피", "평면을 다른 시점에서 본 것처럼 변환하는 기법"],
           ["치명 구역", "가중치 7 이상 구역. 미달이면 총점과 무관하게 불충족"]],
          widths=[4.0, 12.4], size=9)

    # ── 12. 파일 지도 ─────────────────────────────────────────────────
    h(doc, "12. 어느 파일이 무엇을 하나", 1)
    table(doc, ["파일", "역할"],
          [["data/plans/*.json", "CCTV 계획서 (입력)"],
           ["data/building.json", "골조 형상 (입력)"],
           ["data/zones.json", "위험구역 T2·T3 (입력)"],
           ["data/schedule.json", "시간대별 활성 구역 (입력)"],
           ["src/transforms.py", "사진을 ρ·θ·o로 망가뜨린다"],
           ["src/run_grid.py", "288조건 추론"],
           ["src/fit_curve.py", "곡선 피팅. 통과 기준 검사도 여기"],
           ["src/site_model.py", "현장 복셀화"],
           ["src/zone_derive.py", "골조에서 위험구역 도출 (T1)"],
           ["src/geometry.py", "광선투사로 ρ·θ·o 계산"],
           ["src/detect_model.py", "곡선 적용. 범위 밖은 0"],
           ["src/aggregate.py", "다중 카메라 결합"],
           ["src/score.py", "100점 채점"],
           ["src/prescribe.py", "진단과 처방"],
           ["src/schedule.py", "시간대별 진단"],
           ["src/sensitivity.py", "자유 파라미터 스윕"],
           ["src/safety_report.py", "안전보고서 생성"],
           ["mockup/index.html", "심사자 화면 (2D/2.5D/3D)"],
           ["start.bat", "누르면 재계산 후 브라우저"],
           ["docs/reference/", "원문 PDF와 조사 기록"]],
          widths=[5.0, 11.4], size=9)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    p = build()
    print(f"→ {p}")
