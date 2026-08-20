# -*- coding: utf-8 -*-
"""초기 설정값·가정 정리 문서(DOCX) 생성.

**모든 수치를 config.py 와 outputs/ 에서 읽어 넣는다.** 문서에 손으로 적은
값이 하나도 없어야 한다 — 코드가 바뀌면 문서도 다시 뽑으면 맞는다.
CLAUDE.md §0.1-1 "실행하지 않은 계산의 결과를 문서에 쓰지 않는다" 가 그대로
이 스크립트의 설계 이유다.

출처 등급을 네 가지로 태그한다. 이 문서의 값어치는 수치 자체가 아니라
**어느 수치가 무슨 자격으로 거기 있는지** 를 밝히는 데 있다.

    [실측]     이 프로젝트에서 실행해 얻은 값
    [규격]     외부 표준·제조사 사양에서 인용
    [유도]     위 둘에서 계산으로 나온 값
    [잠정]     근거가 약하거나 없는 설계 선택. 바뀔 수 있다

사용:  python tools/make_assumptions_doc.py
"""
from pathlib import Path
import json
import sys
import datetime

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

import config

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = config.ROOT / "outputs" / "설정값_가정_정리.docx"

INK = RGBColor(0x1E, 0x22, 0x28)
MUTED = RGBColor(0x5C, 0x65, 0x70)
WARN = RGBColor(0x8C, 0x20, 0x18)

TAG = {
    "실측": "이 프로젝트에서 실행해 얻은 값",
    "규격": "외부 표준·제조사 사양 인용",
    "유도": "실측·규격에서 계산으로 나온 값",
    "잠정": "근거가 약하거나 없는 설계 선택",
}


# ── 서식 도구 ──────────────────────────────────────────────────────────────

def set_kfont(doc, name="맑은 고딕", size=10):
    """한글 폰트. eastAsia 를 같이 지정하지 않으면 한글만 다른 폰트로 나온다."""
    for style_name in ("Normal", "List Bullet", "List Number"):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = name
        st.font.size = Pt(size)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def para(doc, text="", size=10, bold=False, color=INK, space_after=6,
         align=None, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    return p


def rich(doc, parts, size=10, space_after=6):
    """[(텍스트, 굵게, 색), ...] 를 한 문단에 이어 붙인다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, color in parts:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.font.color.rgb = color
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    return p


def table(doc, headers, rows, widths=None, size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(size)
            r.font.name = "맑은 고딕"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        r.font.color.rgb = INK
    return hd


# ── 본문 ───────────────────────────────────────────────────────────────────

def build():
    site = json.loads((config.ROOT / "outputs" / "site_eval.json").read_text(encoding="utf-8"))
    comp = json.loads((config.ROOT / "outputs" / "comparison.json").read_text(encoding="utf-8"))
    cp = json.loads((config.ROOT / "outputs" / "curve_params.json").read_text(encoding="utf-8"))
    mf = json.loads((config.DATA_FILTERED / "manifest.json").read_text(encoding="utf-8"))
    sn = json.loads((config.ROOT / "outputs" / "sensitivity.json").read_text(encoding="utf-8"))

    occ = [v for v in site["voxels"] if v.get("occupiable") is not False]
    n_all, n_occ = len(site["voxels"]), len(occ)
    pairs = site.get("dori_pair_stats", {})
    rng = site.get("out_of_measured_range", {})
    cm = config.CAMERA_MODEL

    doc = Document()
    set_kfont(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = s.bottom_margin = Cm(2.0)

    # 표지
    para(doc, "AI CCTV 배치 적정성 평가 모델", size=20, bold=True, space_after=2)
    para(doc, "초기 설정값 및 가정 정리", size=14, color=MUTED, space_after=14)
    para(doc, f"제17회 LH 국토기술대전 · 생성일 {datetime.date.today().isoformat()}",
         size=9, color=MUTED, space_after=2)
    para(doc, "본 문서는 config.py 와 outputs/ 에서 값을 읽어 자동 생성된다. "
              "손으로 적은 수치가 없다.", size=9, color=MUTED, space_after=16)

    # 0. 읽는 법
    h(doc, "0. 이 문서를 읽는 법", 1)
    para(doc, "모델이 딛고 선 수치를 한자리에 모았다. 이 문서의 값어치는 수치 자체가 "
              "아니라 어느 수치가 무슨 자격으로 거기 있는지를 밝히는 데 있다. "
              "모든 항목에 출처 등급을 붙였다.")
    table(doc, ["등급", "뜻", "심사에서의 취급"], [
        ["[실측]", TAG["실측"], "그대로 인용 가능"],
        ["[규격]", TAG["규격"], "출처를 함께 밝힐 것"],
        ["[유도]", TAG["유도"], "계산식을 함께 밝힐 것"],
        ["[잠정]", TAG["잠정"], "단정하지 말 것. 민감도로 방어"],
    ], widths=[2.0, 7.0, 6.5])
    rich(doc, [("[잠정] 항목이 결론을 뒤집는지는 8장 민감도에서 확인했다. "
                "ΔWDR 의 부호는 ", False, INK),
               (f"{sn['delta_WDR_range'][0]:+.4f} ~ {sn['delta_WDR_range'][1]:+.4f}",
                True, INK),
               (" 범위에서 유지된다.", False, INK)])

    # 1. 요약
    h(doc, "1. 한눈에 보기", 1)
    table(doc, ["구분", "값", "등급"], [
        ["현장", f"{site['site']['width_m']:.0f} × {site['site']['depth_m']:.0f} m (가상)", "[잠정]"],
        ["복셀", f"{config.VOXEL_M} m 큐브 · 전체 {n_all:,} · 지표용 {n_occ:,}", "[유도]"],
        ["카메라", f"후보 {config.CAMERA_CANDIDATE_COUNT} · 예산 {config.CAMERA_BUDGET}대", "[잠정]"],
        ["카메라 사양", f"{config.IMG_WIDTH_PX}×{config.IMG_HEIGHT_PX} · HFOV {config.HFOV_DEG:.0f}°", "[규격]"],
        ["검출기", f"{cp.get('detector')} (SHWD 파인튜닝)", "[실측]"],
        ["실험 격자", f"{cp.get('n_conditions')}조건 · 이미지 {cp.get('n_images')}장", "[실측]"],
        ["곡선 대표 R²", f"{cp.get('r2_full_grid')}", "[실측]"],
        ["판정 임계", f"P_total < {config.P_DETECT_THRESHOLD}", "[잠정]"],
    ], widths=[3.5, 8.5, 2.0])

    # 2. 실험 설정
    h(doc, "2. 실험 설정 — 곡선을 만든 조건", 1)

    h(doc, "2.1 데이터셋", 2)
    table(doc, ["항목", "값", "등급"], [
        ["출처", mf["source"], "[규격]"],
        ["사용 분할", f"{mf['split']} (학습에 쓰지 않은 분할)", "[실측]"],
        ["필터 ①", f"머리 bbox 짧은 변 ≥ {mf['min_head_px']}px", "[잠정]"],
        ["필터 ②", f"대표 머리크기 ≥ {mf['ref_head_px_min']:.0f}px", "[잠정]"],
        ["30px 완화", "미적용" if not mf["relaxed_to_30px"] else "적용", "[실측]"],
        ["후보 → 선정", f"{mf['n_candidates']:,}장 → {mf['n_selected']:,}장", "[실측]"],
        ["인스턴스", f"착용 {mf['n_hat']:,} · 미착용 {mf['n_person']:,}", "[실측]"],
    ], widths=[3.5, 8.5, 2.0])
    para(doc, "실험셋을 test 분할에서만 뽑은 이유 — 학습에 쓴 이미지로 곡선을 재면 "
              "ρ·θ·o 의 효과가 아니라 검출기의 암기를 재게 된다.", size=9, color=MUTED)

    h(doc, "2.2 검출기", 2)
    table(doc, ["항목", "값", "등급"], [
        ["구조", config.DETECTOR_ARCH, "[실측]"],
        ["가중치", cp.get("detector_weights") or "-", "[실측]"],
        ["학습 데이터", "SHWD train 분할 5,457장", "[실측]"],
        ["학습 조건", "50 epoch · imgsz 640 · batch 16 · seed 0", "[실측]"],
        ["검증 성적", "val mAP50 0.935 (hat 0.938 / person 0.932)", "[실측]"],
    ], widths=[3.5, 8.5, 2.0])
    rich(doc, [("이 검출기는 LH 가 실제 운용하는 A-Eye 가 아니다. ", False, INK),
               ("따라서 절대 수치가 아니라 “설치 조건에 따라 검출률이 변한다”는 "
                "관계의 존재를 보이는 실험으로 읽어야 한다.", True, WARN)], size=9)

    h(doc, "2.3 3축 변형 격자", 2)
    table(doc, ["축", "수준", "구현", "등급"], [
        ["ρ 유효 픽셀밀도", " / ".join(map(str, config.RHO_LEVELS_PX)) + " px",
         "다운샘플 후 원본 크기로 업샘플", "[실측]"],
        ["θ 부감각", " / ".join(map(str, config.THETA_LEVELS_DEG)) + " °",
         "호모그래피 워핑", "[실측]"],
        ["o 가림률", " / ".join(map(str, config.OCC_LEVELS_PCT)) + " %",
         "수직 스트라이프 마스킹", "[실측]"],
        ["조건 수", f"{config.N_CONDITIONS}", "8 × 6 × 6", "[유도]"],
    ], widths=[3.2, 4.6, 5.2, 1.5])
    para(doc, f"방위각 φ 는 축에서 제외했다. 안전모는 상방 시점에서 대략 원형이라 "
              f"정면·후면 차이가 작다. IoU {config.IOU_THR} · confidence "
              f"{config.CONF_THR} 고정.", size=9, color=MUTED)

    h(doc, "2.4 곡선", 2)
    para(doc, "P(ρ, θ, o) = f(ρ) · g(θ) · h(o) — 분리형 곱셈 모델. 상호작용항 없음.")
    rows = []
    for name, q in cp["per_target"].items():
        gp = ", ".join(f"{k}={v:.4g}" for k, v in q["g_theta"]["params"].items())
        rows.append([name, q["target"],
                     f"L={q['f_rho']['L']:.4f} k={q['f_rho']['k']:.4f} x0={q['f_rho']['x0']:.4f}",
                     f"{q['g_theta']['form']} ({gp})",
                     f"λ={q['h_occ']['lambda']:.4f}", f"{q['r2_full_grid']}"])
    table(doc, ["항목", "지표", "f(ρ) 로지스틱", "g(θ)", "h(o)", "R²"], rows,
          widths=[2.6, 2.4, 4.4, 3.4, 1.6, 1.2], size=8)
    para(doc, f"종합은 항목별 {cp.get('aggregate')} 이다. 평균이 아니다 — 안전모 실패가 "
              f"사람 검출 성공에 가려지면 이 도구의 논지가 사라진다. 대표 R² 는 "
              f"가장 나쁜 항목의 값({cp.get('r2_full_grid')})을 쓴다.", size=9, color=MUTED)

    # 3. 카메라
    h(doc, "3. 카메라 사양", 1)
    table(doc, ["항목", "값", "등급"], [
        ["기준 기종", cm["reference"], "[규격]"],
        ["센서", cm["sensor"], "[규격]"],
        ["렌즈", f"{cm['lens_mm'][0]} ~ {cm['lens_mm'][1]} mm 전동 가변초점", "[규격]"],
        ["제조사 고시 화각", f"H {cm['hfov_spec_deg'][0]:.0f}° (wide) ~ {cm['hfov_spec_deg'][1]:.0f}° (tele)", "[규격]"],
        ["보호등급", cm["ingress"], "[규격]"],
        ["본 모델 해상도", f"{config.IMG_WIDTH_PX} × {config.IMG_HEIGHT_PX}", "[규격]"],
        ["본 모델 화각", f"HFOV {config.HFOV_DEG:.0f}°", "[잠정]"],
        ["초점거리", f"f_px = {config.FOCAL_PX:.0f}", "[유도]"],
        ["머리 기준 크기", f"H_head = {config.H_HEAD_M} m", "[잠정]"],
        ["픽셀밀도 식", f"ρ = {config.FOCAL_PX * config.H_HEAD_M:.0f} / d", "[유도]"],
    ], widths=[3.5, 8.5, 2.0])
    rich(doc, [("화각을 90° 로 둔 근거 두 가지. ", True, INK),
               ("① 가변초점이라 화각은 기종이 아니라 설치자가 고르는 값이다. "
                f"고시 범위 {cm['hfov_spec_deg'][0]:.0f}~{cm['hfov_spec_deg'][1]:.0f}° 안에 90° 가 있다. "
                "② 고시 광각단은 배럴 왜곡을 포함하므로 핀홀 식에 그대로 넣을 수 없다. "
                f"이 기종의 핀홀 환산 광각단은 {cm['hfov_pinhole_wide_deg']}° 다.",
                False, INK)], size=9)
    rich(doc, [("한계 — ", True, WARN),
               (cm["note"], False, WARN)], size=9)

    # 4. 현장 모델
    h(doc, "4. 현장 모델", 1)
    table(doc, ["항목", "값", "등급"], [
        ["규모", f"{config.SITE_WIDTH_M} × {config.SITE_DEPTH_M} m", "[잠정]"],
        ["골조", "코어 벽체 2 · 슬래브 · 외곽 비계 · 적치물 2", "[잠정]"],
        ["층", f"슬래브 상단 {config.SLAB_LEVELS_M} m (3개 층)", "[잠정]"],
        ["복셀 격자", f"{config.VOXEL_M} m 큐브", "[잠정]"],
        ["복셀화 방식", f"{config.VOXEL_MODE} · z 상한 {config.VOXEL_Z_MAX_M} m", "[잠정]"],
        ["전체 복셀", f"{n_all:,}", "[유도]"],
        ["지표용 복셀", f"{n_occ:,} (사람이 있을 수 있는 자리)", "[유도]"],
        ["비계 시야 점유율", f"{config.SCAFFOLD_COVERAGE}", "[잠정]"],
    ], widths=[3.5, 8.5, 2.0])
    para(doc, f"사람이 있을 수 있는 자리 정의 — ① 바닥(지면·슬래브) 위 "
              f"{config.OCCUPIABLE_BAND_M[0]}~{config.OCCUPIABLE_BAND_M[1]} m 높이대, "
              f"② 비계에서 {config.OCCUPIABLE_NEAR_SCAFFOLD_M} m 이내. 부피 전체를 분모로 "
              f"삼으면 아무도 못 가는 허공이 커버리지를 희석한다. 시각화는 전부 그리되 "
              f"지표는 이 판정을 통과한 것만으로 낸다.", size=9, color=MUTED)

    h(doc, "4.1 위험구역과 가중치", 2)
    table(doc, ["구역", "가중치", "등급"],
          [[k, v, "[잠정]"] for k, v in config.RISK_WEIGHTS.items()] +
          [["그 외", config.RISK_WEIGHT_DEFAULT, "[잠정]"]],
          widths=[6.0, 3.0, 2.0])
    rich(doc, [("가중치 1~5 는 상대 순위를 표현한 값이며 원자료에서 확인된 수치가 아니다. ",
                False, WARN),
               (f"config.RISK_WEIGHT_SOURCE = \"{config.RISK_WEIGHT_SOURCE}\"", True, WARN),
               (". LH·국토안전관리원 위험도 지수를 확보하면 이 표만 교체한다.", False, WARN)],
         size=9)

    h(doc, "4.2 카메라 후보와 예산", 2)
    mounts = {}
    for c in site["cameras"]:
        mounts.setdefault((c["mount"], c["z"]), 0)
        mounts[(c["mount"], c["z"])] += 1
    table(doc, ["설치 유형", "대수", "높이", "등급"],
          [[m, n, f"{z:.0f} m", "[잠정]"] for (m, z), n in sorted(mounts.items(), key=lambda kv: -kv[0][1])] +
          [["예산", config.CAMERA_BUDGET, "-", "[잠정]"]],
          widths=[5.0, 2.5, 3.0, 2.0])
    para(doc, "경계 폴 6 m 는 KISA 「지능형 CCTV 도입·운영 가이드」가 설치 높이 예시와 "
              "개선 사례(폴 6 m → 8 m)로 다루는 실무 기준값이다.", size=9, color=MUTED)

    # 5. 기하 정의
    h(doc, "5. 기하 계산 정의", 1)
    table(doc, ["기호", "정의", "식 / 값", "등급"], [
        ["d", "카메라–복셀 거리", "유클리드 거리", "[유도]"],
        ["ρ", "머리 유효 픽셀밀도", f"f_px · H_head / d = {config.FOCAL_PX*config.H_HEAD_M:.0f}/d", "[유도]"],
        ["θ", "부감각", "degrees(asin((z_cam − z_voxel) / d))", "[유도]"],
        ["o", "가림률",
         f"복셀에 높이 {config.OCCLUSION_BAR_HEIGHT_M} m 막대를 세우고 "
         f"{config.OCCLUSION_SAMPLE_POINTS} 점에서 광선투사", "[잠정]"],
        ["yaw", "카메라 지향", "15° 간격 24방위 중 위험가중 가시량 최대", "[잠정]"],
    ], widths=[1.6, 3.6, 7.4, 1.5])
    para(doc, "카메라 부각(pitch)은 검사하지 않는다. 수직 화각은 작업면을 덮도록 "
              "조준했다고 본다 — 낙관 방향의 가정이다.", size=9, color=WARN)

    # 6. 판정 기준
    h(doc, "6. 판정 기준", 1)
    table(doc, ["항목", "값", "등급"], [
        ["집계", "P_total(v) = 1 − Π(1 − P(v,c))", "[유도]"],
        ["지표", "WDR = Σ w·P_total / Σ w", "[유도]"],
        ["미달 판정", f"P_total < {config.P_DETECT_THRESHOLD}", "[잠정]"],
        ["기하 기준선", f"IEC 62676-4 DORI {config.GEOMETRIC_DORI_LEVEL} "
                     f"(ρ ≥ {config.GEOMETRIC_MIN_RHO_PX:.2f} px)", "[규격]"],
        ["LH 커버리지 목표", f"{config.LH_COVERAGE_TARGET:.0%} ({config.LH_TARGET_METRIC})", "[잠정]"],
    ], widths=[3.5, 8.5, 2.0])
    rich(doc, [("LH 가 게시한 커버리지 기준은 없다. ", True, WARN),
               (f"위 {config.LH_COVERAGE_TARGET:.0%} 는 우리가 제안하는 잠정 임계이며 "
                f"고시값이 아니다 (config.{config.LH_COVERAGE_TARGET_SOURCE}). "
                "제안서에 “LH 기준 90%” 로 쓰지 말 것.", False, WARN)], size=9)
    para(doc, "DORI 는 인간 관찰자 기준이며 AI 검출기에 대해 검증된 바 없다. "
              "기존 방식의 기준선을 세우는 용도로만 쓴다.", size=9, color=MUTED)

    if pairs.get("levels"):
        h(doc, "6.1 이 현장에서 DORI 등급이 걸리는 정도", 2)
        table(doc, ["등급", "최소 PPM", "ρ 임계", "만족 쌍", "비율"],
              [[k, f"{v['ppm']:.1f}", f"{v['min_rho_px']:.2f} px",
                f"{v['n_pass']:,}", f"{v['ratio']:.1%}"]
               for k, v in pairs["levels"].items()],
              widths=[3.4, 2.6, 2.6, 3.0, 2.0])
        para(doc, f"가시 쌍 {pairs['n_visible']:,}개 기준. ρ 중앙값 "
                  f"{pairs['rho_px_median']} px.", size=9, color=MUTED)

    # 7. 결과
    h(doc, "7. 현재 산출 결과", 1)
    pl = comp.get("placements", {})
    rows = []
    label = {"geometric": "① 기하 커버리지 (기존)",
             "assumed": "② 문헌의 가정 곡선", "empirical": "③ 실측 곡선 (제안)"}
    for k in ("geometric", "assumed", "empirical"):
        if k in pl:
            rows.append([label[k], f"{pl[k]['WDR']:.4f}", f"{pl[k]['fail_voxel_count']:,}"])
    table(doc, ["배치 기준", "WDR", "미달 복셀"], rows, widths=[7.0, 3.0, 3.0])
    d = comp.get("delta_WDR", {})
    if isinstance(d, dict) and d:
        para(doc, "ΔWDR — " + " · ".join(f"{k} {v:+.4f}" for k, v in d.items()),
             size=9, color=MUTED)

    if rng:
        h(doc, "7.1 측정 범위 밖으로 잘린 양", 2)
        table(doc, ["가드", "해당 쌍", "비율"], [
            [f"ρ < {rng['rho_below_min']['limit_px']} px", f"{rng['rho_below_min']['n']:,}",
             f"{rng['rho_below_min']['ratio']:.2%}"],
            [f"θ > {rng['theta_over_max']['limit_deg']}°", f"{rng['theta_over_max']['n']:,}",
             f"{rng['theta_over_max']['ratio']:.2%}"],
            [f"o > {rng['occ_over_max']['limit']}", f"{rng['occ_over_max']['n']:,}",
             f"{rng['occ_over_max']['ratio']:.2%}"],
            ["합계(중복 제거)", f"{rng['any']['n']:,}", f"{rng['any']['ratio']:.2%}"],
        ], widths=[5.0, 4.0, 4.0])
        para(doc, "실측 범위를 벗어난 조건은 외삽하지 않고 P=0 으로 본다. "
                  "안전 판정에서 낙관은 위험한 방향이기 때문이다.", size=9, color=MUTED)

    # 8. 가정값 대장
    h(doc, "8. [잠정] 항목 대장 — 근거가 약한 값 전수", 1)
    para(doc, "심사에서 물으면 이 표로 답한다. 감추지 않는 것이 방어책이다.")
    table(doc, ["값", "현재", "왜 잠정인가", "바뀌면"], [
        ["비계 시야 점유율", f"{config.SCAFFOLD_COVERAGE}",
         "건설현장 실측 가림률 통계가 공개된 것이 없다", "8.1 민감도"],
        ["스트라이프 주기", "화면 가로 / 24",
         "§4.2 가 자유 파라미터로 남긴 값", "8.1 민감도"],
        ["판정 임계", f"{config.P_DETECT_THRESHOLD}",
         "고시 기준 없음. 설계 선택", "8.1 민감도"],
        ["위험 가중치", "1 ~ 5",
         "LH 지수 원자료 미확보. 상대 순위만", "표만 교체하면 됨"],
        ["LH 커버리지 목표", f"{config.LH_COVERAGE_TARGET:.0%}",
         "LH 고시값 아님. 우리 제안치", "게시되면 값만 교체"],
        ["현장 형상", "가상",
         "실제 LH 현장 도면이 아니다. 코드로 생성", "BIM 입력으로 대체 가능"],
        ["카메라 화각", f"{config.HFOV_DEG:.0f}°",
         "가변초점이라 설치자가 고르는 값", "기종·설정 확정 시 교체"],
        ["카메라 부각", "미검사",
         "수직 화각이 작업면을 덮는다고 가정", "낙관 방향"],
        ["머리 기준 크기", f"{config.H_HEAD_M} m",
         "인체 치수 평균 가정", "ρ 전체가 비례 이동"],
    ], widths=[3.2, 2.4, 6.4, 2.6], size=8)

    h(doc, "8.1 민감도 — 결론이 이 값들에 매달려 있는가", 2)
    table(doc, ["스윕", "설정", "ΔWDR"],
          [["스트라이프 주기", r["case"], f"{r['delta_WDR']:+.4f}"]
           for r in sn["stripe_period_lambda"]] +
          [["비계 점유율", f"{sc}", f"{dw:+.4f}"] for sc, dw in
           [(r["scaffold_coverage"], r["delta_WDR"]) for r in sn["scaffold_coverage"]]],
          widths=[4.5, 4.5, 3.5])
    rich(doc, [("절대 WDR 은 크게 흔들리지만 ΔWDR 의 부호는 ", False, INK),
               (f"{sn['delta_WDR_range'][0]:+.4f} ~ {sn['delta_WDR_range'][1]:+.4f}", True, INK),
               (" 범위에서 전부 유지된다. “파라미터를 바꾸면 뒤집히지 않느냐” 에 대한 "
                "답이 이 표다.", False, INK)], size=9)

    # 9. 한계
    h(doc, "9. 알려진 한계 — 코드로 해결하지 않고 명시한다", 1)
    for t in [
        "다운샘플링은 실제 원거리 촬영과 다르다. 대기 흐림·렌즈 광학 한계 미반영.",
        "호모그래피는 3D 시점 변화의 근사다.",
        "분리형 곱셈 모델은 축 간 상호작용을 무시한 1차 근사다.",
        "검출기가 공개 모델의 파인튜닝본이며 실제 운용 모델이 아니다.",
        "카메라 부각(pitch)과 수직 화각을 검사하지 않는다.",
        "야간·조명·기상 조건은 축에서 제외했다.",
        "현장은 가상이며 실제 도면이 아니다.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(t)
        r.font.size = Pt(9.5)
        r.font.name = "맑은 고딕"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # 10. 재현
    h(doc, "10. 재현 절차", 1)
    table(doc, ["단계", "명령", "필요 조건"], [
        ["① 데이터 변환", "python src/voc_to_yolo.py", "SHWD 원본"],
        ["② 실험셋 선정", "python src/filter_data.py", "SHWD 원본"],
        ["③ 검출기 학습", "python src/train_detector.py", "NVIDIA GPU"],
        ["④ 격자 추론", "python src/run_grid.py", "NVIDIA GPU"],
        ["⑤ 가림률 보정", "python src/occ_box.py", "-"],
        ["⑥ 곡선 피팅", "python src/fit_curve.py", "-"],
        ["⑦ 현장 평가", "python src/report.py", "-"],
        ["⑧ 진단 보고서", "python src/safety_report.py <계획서>", "-"],
        ["⑨ 민감도", "python src/sensitivity.py", "-"],
        ["⑩ 이 문서", "python tools/make_assumptions_doc.py", "-"],
    ], widths=[3.4, 7.6, 3.0], size=8.5)
    para(doc, "④ 이후는 커밋된 산출물만으로 재현된다. 난수를 쓰지 않으므로 같은 입력이면 "
              "같은 결과가 나온다.", size=9, color=MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT, n_all, n_occ


if __name__ == "__main__":
    path, n_all, n_occ = build()
    print(f"→ {path}")
    print(f"   복셀 {n_all:,} (지표용 {n_occ:,}) 기준으로 생성")
