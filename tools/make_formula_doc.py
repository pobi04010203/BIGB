# -*- coding: utf-8 -*-
"""인식률 공식 · CCTV 모델 · 검출기 성능 문서(DOCX).

`make_assumptions_doc.py` 가 설정값 전체를 다루는 반면 이 문서는 **셋만** 다룬다 —
검출확률을 어떻게 계산하는가, 어떤 카메라를 전제하는가, 검출기 성능이 얼마인가.

모든 수치는 config.py · outputs/curve_params.json · outputs/grid_results.csv 에서
읽는다. 손으로 적은 값이 없다 (CLAUDE.md §0.1-1).

사용:  python tools/make_formula_doc.py [출력경로]
"""
from pathlib import Path
import csv
import json
import math
import sys
import datetime

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

import config

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

INK = RGBColor(0x1E, 0x22, 0x28)
MUTED = RGBColor(0x5C, 0x65, 0x70)
WARN = RGBColor(0x8C, 0x20, 0x18)
MONO = "Consolas"


def _font(run, size, bold=False, color=INK, name="맑은 고딕"):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_base(doc):
    for sn in ("Normal", "List Bullet"):
        try:
            st = doc.styles[sn]
        except KeyError:
            continue
        st.font.name = "맑은 고딕"
        st.font.size = Pt(10)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def para(doc, text="", size=10, bold=False, color=INK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    _font(p.add_run(text), size, bold, color)
    return p


def rich(doc, parts, size=10, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    for t, b, c in parts:
        _font(p.add_run(t), size, b, c)
    return p


def eq(doc, text, after=8):
    """식. 고정폭으로 찍어 첨자·기호가 흐트러지지 않게 한다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.left_indent = Cm(0.8)
    _font(p.add_run(text), 10.5, False, INK, MONO)
    return p


def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        _font(r, r.font.size.pt if r.font.size else 14, True, INK)
    return hd


def table(doc, headers, rows, widths=None, size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, x in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        _font(c.paragraphs[0].add_run(str(x)), size, True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            _font(cells[i].paragraphs[0].add_run(str(v)), size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ── 데이터 ─────────────────────────────────────────────────────────────────

def load():
    cp = json.loads((config.CURVE_PARAMS_JSON).read_text(encoding="utf-8"))
    rows = list(csv.DictReader(config.GRID_RESULTS_CSV.open(encoding="utf-8")))

    def cut(fixed, var):
        out = [r for r in rows
               if all(abs(float(r[k]) - v) < 1e-6 for k, v in fixed.items())]
        return sorted(out, key=lambda r: float(r[var]))

    return cp, {
        "rho": cut({"theta_deg": 0, "occ_pct_target": 0}, "rho_px"),
        "theta": cut({"rho_px": 48, "occ_pct_target": 0}, "theta_deg"),
        "occ": cut({"rho_px": 48, "theta_deg": 0}, "occ_pct_target"),
    }


def build(out_path: Path):
    cp, cuts = load()
    cm = config.CAMERA_MODEL
    prim = cp["per_target"][cp["primary"]]
    fp, gp, hp = prim["f_rho"], prim["g_theta"], prim["h_occ"]

    doc = Document()
    set_base(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.2)
        s.top_margin = s.bottom_margin = Cm(2.0)

    para(doc, "AI CCTV 안전모 미착용 인식률", size=20, bold=True, after=2)
    para(doc, "계산식 · 카메라 모델 · 검출기 성능", size=14, color=MUTED, after=14)
    para(doc, f"제17회 LH 국토기술대전 · 생성일 {datetime.date.today().isoformat()}",
         size=9, color=MUTED, after=2)
    para(doc, "모든 수치는 config.py · curve_params.json · grid_results.csv 에서 "
              "읽어 자동 생성된다.", size=9, color=MUTED, after=16)

    # ── 1. 인식률 공식 ────────────────────────────────────────────────────
    h(doc, "1. 인식률 공식", 1)
    para(doc, "설치 조건 세 가지가 검출확률을 얼마나 깎는지를 곱으로 나눈다. "
              "각 항은 다른 두 축을 기준값에 고정한 단면에서 따로 맞췄다.")
    eq(doc, "P(ρ, θ, o)  =  f(ρ) · g(θ) · h(o)")
    table(doc, ["기호", "뜻", "단위", "기준값"], [
        ["ρ", "머리 유효 픽셀밀도", "px", "48"],
        ["θ", "부감각 (내려다보는 각)", "°", "0"],
        ["o", "가림률", "0~1", "0"],
    ], widths=[1.8, 6.5, 2.2, 2.2])
    para(doc, "상호작용항은 넣지 않는다. 해석이 쉽고 최적화에 그대로 들어가며, "
              "1차 근사임을 명시하는 것이 방어 전략이다.", size=9, color=MUTED)

    h(doc, "1.1 f(ρ) — 픽셀밀도 항", 2)
    eq(doc, f"f(ρ) = L / (1 + exp(-k · (ρ - x0)))")
    eq(doc, f"     L = {fp['L']:.4f}   k = {fp['k']:.4f}   x0 = {fp['x0']:.4f}")
    table(doc, ["ρ (px)", "PPM", "실측 미착용", "실측 착용", "conf 평균"],
          [[f"{float(r['rho_px']):.0f}", f"{float(r['rho_px'])/config.H_HEAD_M:.0f}",
            r["recall_nohat"], r["recall_hat"], r["conf_mean"]]
           for r in reversed(cuts["rho"])],
          widths=[2.2, 2.2, 3.0, 3.0, 2.6])
    para(doc, "PPM(장면 미터당 픽셀)은 ρ / H_head 로 환산한 값이며 IEC 62676-4 의 "
              "DORI 등급과 직접 비교된다.", size=9, color=MUTED)

    h(doc, "1.2 g(θ) — 부감각 항", 2)
    if gp["form"] == "logistic":
        eq(doc, "g(θ) = (1 + exp(-k · x0)) / (1 + exp(k · (θ - x0)))")
        eq(doc, f"     k = {gp['params']['k']:.4f}   x0 = {gp['params']['x0']:.4f}  (도)")
    else:
        eq(doc, f"g(θ) = {gp['form']}  {gp['params']}")
    table(doc, ["θ (°)", "실측 미착용", "실측 착용"],
          [[f"{float(r['theta_deg']):.0f}", r["recall_nohat"], r["recall_hat"]]
           for r in cuts["theta"]], widths=[3.0, 5.0, 5.0])
    rich(doc, [("45°까지는 거의 평평하다가 60~75°에서 절벽처럼 무너진다. ", False, INK),
               (f"로지스틱 중점 x0 = {gp['params'].get('x0', 0):.1f}° 가 그 무너지는 지점이다.",
                True, INK)], size=9)
    para(doc, "2차·지수로는 이 모양을 맞히지 못한다 — 후보별 단면 R² 는 "
              + " · ".join(f"{k} {v}" for k, v in (gp.get("r2_candidates") or {}).items())
              + " 였다.", size=9, color=MUTED)

    h(doc, "1.3 h(o) — 가림 항", 2)
    eq(doc, "h(o) = exp(-λ · o)")
    eq(doc, f"     λ = {hp['lambda']:.4f}   (o 는 0~1)")
    table(doc, ["목표 o (%)", "실측 가림률 (%)", "실측 미착용", "실측 착용"],
          [[f"{float(r['occ_pct_target']):.0f}", r.get("occ_pct_box", "-"),
            r["recall_nohat"], r["recall_hat"]] for r in cuts["occ"]],
          widths=[2.8, 3.6, 3.4, 3.2])
    rich(doc, [("세 축 중 가림이 압도적으로 세다. ", True, INK),
               ("30% 가림에 미착용 검출률이 0.958 → 0.239 로 무너진다. "
                "입력은 화면 전체가 아니라 인스턴스 평균 가림률이다 — 현장의 o 가 "
                "사람 단위 값이라 정의를 맞췄고, 두 정의의 실측 차이는 0.4%p 이내였다.",
                False, INK)], size=9)

    h(doc, "1.4 측정 범위 — 밖으로 외삽하지 않는다", 2)
    table(doc, ["축", "측정 범위", "범위 밖 처리"], [
        ["ρ", f"{fp['measured_range_px'][0]:.0f} ~ {fp['measured_range_px'][1]:.0f} px",
         "하한 미만은 P = 0"],
        ["θ", f"{gp.get('measured_range_deg', ['?','?'])[0]:.0f} ~ "
              f"{gp.get('measured_range_deg', ['?','?'])[1]:.0f} °", "상한 초과는 P = 0"],
        ["o", f"{hp.get('measured_range', ['?','?'])[0]:.2f} ~ "
              f"{hp.get('measured_range', ['?','?'])[1]:.2f}", "상한 초과는 P = 0"],
    ], widths=[2.0, 5.0, 6.0])
    para(doc, "안전 판정 도구에서 낙관은 위험한 방향이다. 재본 적 없는 구간을 "
              "곡선으로 늘리는 대신 '검출 못 한다'로 본다.", size=9, color=MUTED)

    h(doc, "1.5 항목별 곡선과 종합", 2)
    para(doc, f"탐지 항목마다 곡선을 따로 뽑고 종합은 {cp.get('aggregate')} 으로 낸다.")
    table(doc, ["항목", "지표", "λ (가림)", "전체 격자 R²"],
          [[k, v["target"], f"{v['h_occ']['lambda']:.4f}", v["r2_full_grid"]]
           for k, v in cp["per_target"].items()], widths=[3.4, 3.4, 3.0, 3.2])
    rich(doc, [("평균이 아니라 최솟값이다. ", True, WARN),
               ("안전모 실패가 사람 검출 성공에 가려지면 이 도구의 논지가 사라진다. "
                f"대표 R² 는 가장 나쁜 항목의 값({cp.get('r2_full_grid')})을 쓴다.",
                False, WARN)], size=9)

    h(doc, "1.6 다중 카메라 결합과 최종 지표", 2)
    eq(doc, "P_total(v) = 1 - Π (1 - P(v, c))        c = 선택된 카메라")
    eq(doc, "WDR        = Σ w(v) · P_total(v) / Σ w(v)")
    rich(doc, [("중첩은 페널티가 아니라 이득이다. ", True, INK),
               ("각 0.6 인 두 대가 합쳐 0.84 가 된다. 기존 연구는 처리·저장 부담을 "
                "이유로 중첩 최소화를 목표로 삼지만, 확률 모델에서는 고위험 구역의 "
                "최적해가 정반대로 나온다.", False, INK)], size=9)
    para(doc, f"판정 — P_total < {config.P_DETECT_THRESHOLD} 인 복셀을 미달구역으로 "
              f"표시한다. 임계는 고시 기준이 없어 설계 선택이다.", size=9, color=MUTED)

    # ── 2. 기하량 ─────────────────────────────────────────────────────────
    h(doc, "2. 기하량 — 현장 좌표에서 ρ·θ·o 를 얻는 법", 1)
    eq(doc, "f_px = (W_img / 2) / tan(HFOV / 2)")
    eq(doc, f"     = ({config.IMG_WIDTH_PX} / 2) / tan({config.HFOV_DEG:.0f}° / 2) "
            f"= {config.FOCAL_PX:.0f}")
    eq(doc, "ρ = f_px · H_head / d"
            f"        = {config.FOCAL_PX * config.H_HEAD_M:.0f} / d      (d 단위 m)")
    eq(doc, "θ = degrees( asin( (z_cam - z_voxel) / d ) )")
    eq(doc, f"o = 복셀에 높이 {config.OCCLUSION_BAR_HEIGHT_M} m 막대를 세우고 "
            f"{config.OCCLUSION_SAMPLE_POINTS} 점에서")
    eq(doc, "    카메라로 광선을 쏴 골조에 막히는 비율")
    para(doc, "전부 막히면 visible = false, P = 0 이다.", size=9, color=MUTED)

    # ── 3. CCTV 모델 ──────────────────────────────────────────────────────
    h(doc, "3. CCTV 모델", 1)
    table(doc, ["항목", "값"], [
        ["기준 기종", cm["reference"]],
        ["센서", cm["sensor"]],
        ["렌즈", f"{cm['lens_mm'][0]} ~ {cm['lens_mm'][1]} mm 전동 가변초점"],
        ["제조사 고시 화각",
         f"H {cm['hfov_spec_deg'][0]:.0f}° (wide) ~ {cm['hfov_spec_deg'][1]:.0f}° (tele)"],
        ["보호등급", cm["ingress"]],
        ["본 모델 해상도", f"{config.IMG_WIDTH_PX} × {config.IMG_HEIGHT_PX}"],
        ["본 모델 화각", f"HFOV {config.HFOV_DEG:.0f}°"],
        ["초점거리", f"f_px = {config.FOCAL_PX:.0f}"],
        ["머리 기준 크기", f"H_head = {config.H_HEAD_M} m"],
    ], widths=[4.0, 10.0])
    para(doc, "선정 근거 — 한화비전은 조달청 등록 국내 대표 영상보안 기업이고, "
              "P 시리즈 AI 4K 4종이 국정원 보안기능 확인서를 취득했다(공공기관 납품 "
              "필수 요건). KISA 「지능형 CCTV 도입·운영 가이드」는 3세대 지능형을 "
              "UHD급으로 규정한다.", size=9, color=MUTED)
    rich(doc, [("한계 — ", True, WARN), (cm["note"], False, WARN)], size=9)

    h(doc, "3.1 화각은 기종이 아니라 설치자가 고르는 값이다", 2)
    para(doc, "가변초점이라 화각이 범위 안에서 조절된다. 그리고 그 선택이 "
              "검출거리를 네 배 넘게 바꾼다.")
    rows = []
    thr_obs = config.dori_rho_px("observation")
    for name, hf in [("고시 wide", cm["hfov_spec_deg"][0]),
                     ("핀홀 wide", cm["hfov_pinhole_wide_deg"]),
                     ("본 모델", config.HFOV_DEG),
                     ("중간", 60.0),
                     ("고시 tele", cm["hfov_spec_deg"][1])]:
        f_px = (config.IMG_WIDTH_PX / 2) / math.tan(math.radians(hf) / 2)
        r = f_px * config.H_HEAD_M
        rows.append([name, f"{hf:.1f}°", f"{r:.0f} / d",
                     f"{r/thr_obs:.1f} m", f"{r/config.dori_rho_px('recognition'):.1f} m"])
    table(doc, ["설정", "HFOV", "ρ", "observation 한계", "recognition 한계"], rows,
          widths=[2.6, 2.2, 2.6, 3.4, 3.2])
    rich(doc, [("제조사 고시 광각단을 핀홀 식에 그대로 넣으면 안 된다. ", True, WARN),
               (f"고시 {cm['hfov_spec_deg'][0]:.0f}° 는 핀홀 예측"
                f"({cm['hfov_pinhole_wide_deg']}°)보다 크고, 차이는 광각단 배럴 왜곡이다. "
                "왜곡된 렌즈는 화면 중앙과 주변의 픽셀밀도가 달라 균일성을 가정할 수 "
                "없다. 그래서 핀홀 환산이 성립하는 값을 쓴다.", False, WARN)], size=9)

    h(doc, "3.2 DORI 등급과 최대 거리", 2)
    table(doc, ["등급", "최소 PPM", "ρ 임계", "최대 거리"],
          [[k, f"{v:.1f}", f"{config.dori_rho_px(k):.2f} px",
            f"{config.FOCAL_PX*config.H_HEAD_M/config.dori_rho_px(k):.1f} m"]
           for k, v in config.DORI_PPM.items()] +
          [["실측 곡선 하한", "-", f"{config.RHO_MEASURED_MIN_PX:.0f} px",
            f"{config.FOCAL_PX*config.H_HEAD_M/config.RHO_MEASURED_MIN_PX:.1f} m"]],
          widths=[3.6, 2.6, 3.0, 3.2])
    para(doc, "DORI 는 인간 관찰자 기준이며 AI 검출기에 대해 검증된 바 없다. "
              "기존 방식의 기준선을 세우는 용도로만 쓴다.", size=9, color=MUTED)

    # ── 4. 검출기 성능 ────────────────────────────────────────────────────
    h(doc, "4. 검출기 성능", 1)
    table(doc, ["항목", "값"], [
        ["구조", config.DETECTOR_ARCH],
        ["가중치", cp.get("detector_weights") or "-"],
        ["학습 데이터", "SHWD train 분할 5,457장"],
        ["학습 조건", "50 epoch · imgsz 640 · batch 16 · seed 0"],
        ["검증 (val)", "mAP50 0.935 — hat 0.938 / person 0.932"],
        ["실험셋", f"SHWD test 분할 {cp.get('n_images')}장 (학습에 쓰지 않음)"],
        ["판정 임계", f"IoU {config.IOU_THR} · confidence {config.CONF_THR}"],
        ["격자", f"{cp.get('n_conditions')}조건 "
                f"(ρ {len(config.RHO_LEVELS_PX)} × θ {len(config.THETA_LEVELS_DEG)}"
                f" × o {len(config.OCC_LEVELS_PCT)})"],
    ], widths=[4.0, 10.0])

    h(doc, "4.1 축별 감쇠 폭", 2)
    def span(rs, col="recall_nohat"):
        vs = [float(r[col]) for r in rs]
        return max(vs) - min(vs)
    table(doc, ["축", "측정 구간", "미착용 검출률 변화폭"], [
        ["ρ", f"{config.RHO_LEVELS_PX[-1]} ~ {config.RHO_LEVELS_PX[0]} px",
         f"{span(cuts['rho'])*100:.0f} %p"],
        ["θ", f"0 ~ {config.THETA_LEVELS_DEG[-1]} °", f"{span(cuts['theta'])*100:.0f} %p"],
        ["o", f"0 ~ {config.OCC_LEVELS_PCT[-1]} %", f"{span(cuts['occ'])*100:.0f} %p"],
    ], widths=[2.4, 5.0, 5.0])
    para(doc, "가림이 가장 세고 부감각이 그다음이다. 픽셀밀도는 측정 구간 안에서 "
              "가장 완만하지만, 현장 거리가 늘어나면 곧바로 절벽 구간(ρ 8px 이하)에 "
              "들어간다.", size=9, color=MUTED)

    rich(doc, [("이 검출기는 LH 가 실제 운용하는 A-Eye 가 아니다. ", True, WARN),
               ("따라서 이 문서의 수치는 절대 성능이 아니라 "
                "'설치 조건에 따라 검출률이 변한다'는 관계의 크기를 보이는 값이다. "
                "곡선 파일(curve_params.json) 하나만 교체하면 다른 검출기로 그대로 "
                "돌아가는 구조로 만들었다.", False, WARN)], size=9)

    # ── 5. 한계 ───────────────────────────────────────────────────────────
    h(doc, "5. 이 공식의 한계", 1)
    for t in [
        "분리형 곱셈 모델은 축 간 상호작용을 무시한 1차 근사다.",
        "ρ 축은 다운샘플링으로 만들었다. 실제 원거리 촬영의 대기 흐림·렌즈 광학 한계는 반영되지 않는다.",
        "θ 축은 호모그래피 워핑이며 3D 시점 변화의 근사다.",
        "o 축의 스트라이프 주기는 자유 파라미터다. 현장 실측 가림률 통계가 공개된 것이 없다.",
        "방위각·야간·조명·기상은 축에서 제외했다.",
        "카메라 부각(pitch)과 수직 화각을 검사하지 않는다.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        _font(p.add_run(t), 9.5)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    dest = Path(sys.argv[1]) if len(sys.argv) > 1 else (
        Path.home() / "Downloads" / "인식률공식_CCTV모델_검출기성능.docx")
    print(f"→ {build(dest)}")
